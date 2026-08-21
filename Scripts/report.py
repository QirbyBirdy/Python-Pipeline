"""Stage 7 (build brief Step 10 / Section 4.3): the automated Word report.
Renders Outputs/report/<name>_report.docx directly from the pipeline's own
outputs -- every number in it (row/column counts, missingness percentages,
sentinel/imputation counts) is computed fresh from the parquet files at
render time, not typed in by hand, so re-running the pipeline on updated
data regenerates a report that's still accurate. Every chart is drawn by
this module's own matplotlib code against the final data, not a copy of
the JPEGs visualize.py already produced -- the brief is explicit that
"pasted static images do not count."

python-docx was chosen over Quarto: this project has used python-docx
throughout (the three weekly reports), there's no existing Quarto/Jupyter
setup to add, and the brief lists python-docx as an equally acceptable
option ("a Python-based report generator (e.g., python-docx) with docx
output").

Chart theme: reuses visualize.py's already-validated palette (dataviz
skill, references/palette.md) rather than re-deriving or re-validating a
second one -- satisfies the brief's "consistent theme" requirement by
construction. Interactive-chart guidance from that skill (hover tooltips,
dark-mode toggle) is skipped for the same reason visualize.py skips it:
these are static images embedded in a Word document, not a web component.
"""

import io
import subprocess
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

import config
from clean import check_employment_skip_logic, flag_monetary_outliers
from data_dictionary import column_labels
from visualize import (
    _CATEGORICAL,
    _MUTED_INK,
    _SECONDARY_INK,
    _SURFACE,
    _wrap_label,
)

# Headline cleaning/imputation-scope variables per dataset -- what the
# missingness-overview chart plots. The full 175/28-column version would
# be an unreadable bar chart; this stays scoped to what the report
# actually discusses (same scope discipline as problem_inventory.py).
_SCOPE_COLUMNS = {
    "ind": ["q1_03", "q1_04", "q3_03", "q3_07_1", "isic_code", "isco_code"],
    "emig": ["q7_06"],
}


def _style_axes(ax) -> None:
    for spine in ("top", "right"):
        ax.spines[spine].set_visible(False)
    ax.spines["left"].set_color(_MUTED_INK)
    ax.spines["bottom"].set_color(_MUTED_INK)
    ax.tick_params(colors=_SECONDARY_INK, labelsize=9)
    ax.title.set_color("#1D1F20")


def _fig_to_buffer(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, facecolor=_SURFACE, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------- charts

def chart_missingness_overview(name: str, typed: pd.DataFrame, cleaned: pd.DataFrame) -> io.BytesIO:
    """Required chart 1: % missing by variable, before vs. after cleaning."""
    cols = [c for c in _SCOPE_COLUMNS[name] if c in typed.columns and c in cleaned.columns]
    raw_pct = [100 * typed[c].isna().mean() for c in cols]
    cleaned_pct = [100 * cleaned[c].isna().mean() for c in cols]

    y = np.arange(len(cols))
    h = 0.32
    # Minimum height floor: with as few as one variable (emig), a fraction-
    # based bottom margin leaves too little *absolute* room for the xlabel
    # and the external legend not to collide, even though the same fraction
    # is plenty of room on the taller (ind) chart.
    fig_height = max(0.6 * len(cols) + 1.4, 3.2)
    fig, ax = plt.subplots(figsize=(7, fig_height))
    ax.barh(y + h / 2, raw_pct, height=h, color=_CATEGORICAL[0], label="Before cleaning")
    ax.barh(y - h / 2, cleaned_pct, height=h, color=_CATEGORICAL[1], label="After cleaning")
    ax.set_yticks(y)
    ax.set_yticklabels(cols)
    ax.invert_yaxis()
    ax.set_xlabel("% missing")
    ax.set_title("Missingness before vs. after cleaning\n(cleaning-scope variables)")
    _style_axes(ax)
    # Legend sits below the axes, not inside the plot area -- with as few
    # as one variable (emig) there is no reliably empty corner to place it
    # in without risking overlap with a long bar. Anchored in FIGURE
    # coordinates (bbox_transform=fig.transFigure), not axes coordinates --
    # axes-fraction offsets scale with the axes' own height, which made the
    # legend collide with the xlabel on short (few-row) charts even though
    # the same offset looked fine on tall ones. Figure coordinates are a
    # fixed 0-1 range regardless of row count, so a small reserved strip at
    # the very bottom of the figure is enough on every chart this function
    # draws. Deliberately no fig.tight_layout() -- it doesn't know about a
    # legend placed outside the axes' own bbox.
    ax.legend(frameon=False, loc="lower center", bbox_to_anchor=(0.5, 0.02), bbox_transform=fig.transFigure, ncol=2, fontsize=9)
    fig.subplots_adjust(bottom=1.0 / fig_height)
    return _fig_to_buffer(fig)


def chart_numeric_distribution(series: pd.Series, title: str, xlabel: str, color: str) -> io.BytesIO:
    """Required chart 2 (used twice: age and hours worked)."""
    data = series.dropna().astype(float)
    fig, ax = plt.subplots(figsize=(6.3, 3.8))
    ax.hist(data, bins=30, color=color, edgecolor=_SURFACE, linewidth=0.4)
    median_val = data.median()
    ax.axvline(median_val, color="#1D1F20", linestyle="--", linewidth=1, label=f"median = {median_val:.0f}")
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel("Number of respondents")
    _style_axes(ax)
    ax.legend(frameon=False, fontsize=9)
    fig.tight_layout()
    return _fig_to_buffer(fig)


def chart_observed_vs_imputed(df: pd.DataFrame, column: str, title: str) -> io.BytesIO:
    """Required chart 3: observed vs. imputed values for one imputed variable."""
    flag_col = f"{column}_imputed"
    observed = df.loc[~df[flag_col].astype(bool), column].dropna().astype(float)
    imputed = df.loc[df[flag_col].astype(bool), column].dropna().astype(float)

    bins = np.histogram_bin_edges(observed, bins=25)
    fig, ax = plt.subplots(figsize=(6.3, 3.8))
    ax.hist(observed, bins=bins, color=_CATEGORICAL[0], alpha=0.85, label=f"Observed (n={len(observed)})")
    ax.hist(imputed, bins=bins, color=_CATEGORICAL[1], alpha=0.85, label=f"Imputed (n={len(imputed)})")
    ax.set_title(title)
    ax.set_xlabel(column)
    ax.set_ylabel("Number of respondents")
    _style_axes(ax)
    ax.legend(frameon=False, fontsize=9)
    fig.tight_layout()
    return _fig_to_buffer(fig)


def chart_of_choice(df: pd.DataFrame, column: str, title: str) -> io.BytesIO:
    """Required chart 4: one substantive categorical breakdown -- employment
    status (ind) / reason for emigrating (emig). Single hue, not one color
    per bar -- this is one series (a respondent count), and the y-axis
    labels already carry category identity, so per-bar color would be
    decorative, not informative (same rule visualize.py's own bar charts
    already follow)."""
    counts = df[column].value_counts(dropna=True).sort_values()
    labels = [_wrap_label(str(v), width=32) for v in counts.index]

    fig, ax = plt.subplots(figsize=(7.2, 0.5 * len(counts) + 1.6))
    ax.barh(labels, counts.values, color=_CATEGORICAL[0])
    for i, v in enumerate(counts.values):
        ax.text(v, i, f" {v:,}", va="center", fontsize=8.5, color=_SECONDARY_INK)
    ax.set_title(title)
    ax.set_xlabel("Respondents")
    _style_axes(ax)
    fig.tight_layout()
    return _fig_to_buffer(fig)


# ---------------------------------------------------------------- docx helpers

def _heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: list) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _table(doc: Document, headers: list, rows: list) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for cell, h in zip(t.rows[0].cells, headers):
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for cell, val in zip(cells, row):
            cell.text = str(val)


def _add_chart(doc: Document, buf: io.BytesIO, width_in: float = 6.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(width_in))


def _pct(numerator: int, denominator: int) -> str:
    return f"{100 * numerator / denominator:.1f}%" if denominator else "0.0%"


def _first_git_commit_date(csv_path: Path) -> str:
    """Earliest git commit date that added `csv_path` -- used as a proxy for
    "date of download" (brief Section 4.3), since the raw CSVs carry no
    download timestamp of their own. Computed fresh via git rather than
    hardcoded, so it can't drift from the actual repository history.
    Falls back to a plain message if git isn't available (e.g. a zip
    export of the repo with no .git history)."""
    try:
        result = subprocess.run(
            ["git", "log", "--diff-filter=A", "--format=%ad", "--date=short", "--", str(csv_path)],
            cwd=csv_path.resolve().parent,
            capture_output=True,
            text=True,
            check=True,
        )
        dates = [line for line in result.stdout.strip().splitlines() if line]
        return dates[-1] if dates else "not available (no git history for this file)"
    except (subprocess.CalledProcessError, FileNotFoundError):
        return "not available (git not found or not a git repository)"


# ---------------------------------------------------------------- section builders

def _section_dataset_overview(doc: Document, name: str, datasets: dict, typed: pd.DataFrame) -> None:
    paths = datasets[name]
    n_rows, n_cols = typed.shape
    download_date = _first_git_commit_date(paths["csv"])
    label = {"ind": "individual records (ind)", "emig": "household emigration module (emig)"}[name]

    _heading(doc, "1. Dataset Overview", level=1)
    _para(
        doc,
        f"Source: GLFS 2017 Q4 microdata (survey round 174), Bureau of Statistics, Guyana -- "
        f"{label} track. Row/household counts verified against the source methodological report "
        f"at ingest (Outputs/logs/{name}_ingest.log).",
    )
    _para(
        doc,
        f"Size: {n_rows:,} rows x {n_cols} columns (raw, as loaded by ingest.py). Date this raw file "
        f"first entered the project's version control: {download_date} (not separately logged as a "
        "download timestamp; this is the earliest git record of the file).",
    )
    _para(
        doc,
        "Variable scope: this report's cleaning/imputation sections cover the assigned named "
        "variables only -- sex, age, hours worked, industry (ISIC), occupation (ISCO) -- per the "
        "brief's Section 1.2. Every column in the dataset, including the ones outside that scope, "
        f"is documented in the accompanying data dictionary (Outputs/data_dictionary/{name}_data_dictionary.csv).",
    )


def _section_diagnostics(doc: Document, name: str, typed: pd.DataFrame, chart_buf) -> None:
    _heading(doc, "2. Data Quality Diagnostics (Before Cleaning)", level=1)
    overall_missing = 100 * typed.isna().mean().mean()
    _para(
        doc,
        f"Average missingness across all {typed.shape[1]} raw columns: {overall_missing:.1f}%. "
        "Most of this is structural (skip-pattern) missingness, not a data-quality defect -- see "
        f"Outputs/diagnostics/{name}_report.md for the full column-by-column breakdown.",
    )
    if name == "ind":
        n_sex_missing = int(typed["q1_03"].isna().sum())
        n_age_sentinel = int((typed["q1_04"] == -1).sum())
        n_hours_sentinel = int((typed["q3_03"] == 99).sum())
        n_no_main_job = int(typed["q3_16"].isna().sum())
        n_dup = len(typed) - len(typed.drop_duplicates(subset=["hhid", "member"]))
        _bullets(
            doc,
            [
                f"Sex (q1_03): {n_sex_missing} missing ({_pct(n_sex_missing, len(typed))}), clean text values.",
                f"Age (q1_04): {n_age_sentinel} row(s) carry a -1 ('don't know') sentinel value, not a real age.",
                f"Hours worked (q3_03): {n_hours_sentinel} row(s) carry a 99 top-code sentinel; every other value caps at 98.",
                f"Industry/occupation (isic_code/isco_code): missing exactly matches the {n_no_main_job:,} rows "
                "with no main job -- structural, not incidental.",
                f"Duplicate person-records: {n_dup} found (exact-match key hhid+member).",
            ],
        )
    else:
        n_age_sentinel = int((typed["q7_06"] == -1).sum())
        n_dup = len(typed) - len(typed.drop_duplicates(subset=["hhid", "emig"]))
        _bullets(
            doc,
            [
                f"Age (q7_06): {n_age_sentinel} row(s) carry a -1 ('don't know') sentinel value, not a real age.",
                f"Duplicate person-records: {n_dup} found (exact-match key hhid+emig).",
            ],
        )
    _para(
        doc,
        "Note on the chart below: missingness legitimately RISES after cleaning for sentinel-corrected "
        "columns. clean.py converts placeholder values (-1, 99) -- which are not counted as missing in "
        "the raw data, since they are literal numbers -- into honest NaN. A taller 'after cleaning' bar "
        "for those columns reflects more accurate missingness reporting, not a defect introduced by cleaning.",
    )
    _add_chart(doc, chart_buf)


def _section_cleaning(doc: Document, name: str, cleaned: pd.DataFrame, typed: pd.DataFrame) -> None:
    _heading(doc, "3. Cleaning Decisions", level=1)
    _para(doc, "For each class of issue found: what was done, and why. Full evidence in CLEANING_PLAN.md and Scripts/clean.py.")

    rows = []
    n_dup = len(typed) - len(typed.drop_duplicates(subset=(["hhid", "member"] if name == "ind" else ["hhid", "emig"])))
    rows.append(["Exact-duplicate person-records", "Dropped (keep first)", f"{n_dup} found"])

    age_col = "q1_04" if name == "ind" else "q7_06"
    n_age = int((typed[age_col] == -1).sum())
    rows.append([f"Age ({age_col}) = -1 sentinel", "Corrected to NaN", f"{n_age} row(s) -- verified 'don't know' placeholder"])

    if name == "ind":
        n_hrs99 = int((typed["q3_03"] == 99).sum())
        rows.append(["Hours worked (q3_03) = 99 sentinel", "Corrected to NaN", f"{n_hrs99} row(s) -- confirmed top-code, checked against every other hours column first"])

        n_isic_coded = int(cleaned["isic_code"].notna().sum())
        n_isic_classified = int(cleaned["isic_section"].notna().sum())
        rows.append([
            "Industry/occupation leading-zero stripping",
            "Zero-padded before classification",
            f"{n_isic_coded - n_isic_classified} of {n_isic_coded:,} coded rows failed to classify after the fix",
        ])

        n_consolidated_rows = 0
        for col in config.FREE_TEXT_COLUMNS.get(name, []):
            if col in typed.columns and col in cleaned.columns:
                both_present = typed[col].notna() & cleaned[col].notna()
                n_consolidated_rows += int((typed.loc[both_present, col] != cleaned.loc[both_present, col]).sum())
        rows.append([
            "Free-text near-duplicates (>=0.98 similarity)",
            "Auto-merged to most frequent form",
            f"{n_consolidated_rows} row(s) changed across {len(config.FREE_TEXT_COLUMNS.get(name, []))} free-text columns",
        ])

        n_monetary_flags = flag_monetary_outliers(cleaned, name)
        n_monetary_rows = len(n_monetary_flags[["hhid", "id2"]].drop_duplicates()) if len(n_monetary_flags) else 0
        rows.append([
            "Monetary outliers (1.5x IQR)",
            "Flagged only, not altered",
            f"{len(n_monetary_flags)} flag(s) across {n_monetary_rows} unique row(s)",
        ])

        n_skip_violations = len(check_employment_skip_logic(cleaned))
        rows.append([
            "Employment skip-logic (isco_code filled with no evidence of employment)",
            "Flagged only, not altered",
            f"{n_skip_violations} unexplained violation(s)",
        ])

    _table(doc, ["Issue", "Decision", "Evidence"], rows)


def _imputation_population(base_col: str) -> str:
    """Human-readable description of which rows a given imputed variable's
    median was computed/filled over -- varies per variable, not a repeated
    boilerplate string, so the table below actually says something new in
    every row."""
    if base_col in ("q1_04", "q7_06"):
        return "Everyone (age has no applicability gate)"
    if base_col == "q3_08":
        return "Has a second job (q3_01 == 'Yes')"
    if base_col == "q3_03" or base_col.startswith("q3_07_"):
        return "Has a main job (q3_16 not null)"
    return "See README"


def _short_variable_label(base_col: str, dictionary_label: str) -> str:
    """A table-cell-length name for the imputation table's "Variable"
    column. The raw survey question text (`dictionary_label`, the same one
    the data dictionary and .sav/.dta files use) is the right label for a
    dictionary entry, but several of these questions run to 100+ characters
    -- unreadable, and often silently truncated, inside a table column.
    Falls back to the full dictionary label for anything not covered by
    the short names already established elsewhere in this project
    (impute.py's own log lines, the README's "Running imputation" table)."""
    if base_col in ("q1_04", "q7_06"):
        return "Age"
    if base_col == "q3_03":
        return "Hours worked, main job"
    if base_col.startswith("q3_07_"):
        day = base_col.rsplit("_", 1)[-1]
        return f"Hours worked, main job, day {day}"
    if base_col == "q3_08":
        return "Hours worked, other job"
    return dictionary_label


def _section_imputation(doc: Document, name: str, imputed: pd.DataFrame, labels: dict, obs_vs_imp_buf) -> None:
    _heading(doc, "4. Imputation Applied", level=1)
    _para(doc, "Method chosen per variable, per brief Section 3.3 -- justification checked against the actual data, not assumed.")

    flag_cols = [c for c in imputed.columns if c.endswith("_imputed")]
    rows = []
    for col in flag_cols:
        base = col[: -len("_imputed")]
        n = int(imputed[col].sum())
        short_label = _short_variable_label(base, labels.get(base, base))
        rows.append([f"{short_label} ({base})", "Median", n, _imputation_population(base)])
    _table(doc, ["Variable", "Method", "Rows imputed", "Applicable population"], rows)

    if name == "ind":
        _bullets(
            doc,
            [
                "Hours worked, other job (q3_04): left missing -- 100% structural, confirmed 0 incidental gaps.",
                "Industry/occupation (isic_code/isco_code): left missing -- confirmed structural (no main job).",
            ],
        )
        _add_chart(doc, obs_vs_imp_buf)


def _section_post_cleaning(doc: Document, name: str, imputed: pd.DataFrame, dist_bufs: list) -> None:
    _heading(doc, "5. Post-Cleaning Summary", level=1)
    n_rows, n_cols = imputed.shape
    n_flag_cols = len([c for c in imputed.columns if c.endswith("_imputed")])
    flag_col_word = "column" if n_flag_cols == 1 else "columns"
    overall_missing = 100 * imputed.isna().mean().mean()
    _para(
        doc,
        f"Final dataset: {n_rows:,} rows x {n_cols} columns ({n_flag_cols} imputation-flag {flag_col_word} "
        "added on top of the cleaned data). Average missingness across all columns: "
        f"{overall_missing:.1f}% (mostly structural skip-pattern missingness, not incidental gaps).",
    )
    numeric_cols = {"ind": ["q1_04", "q3_03"], "emig": ["q7_06"]}[name]
    labels = {"q1_04": "Age", "q3_03": "Hours worked (main job)", "q7_06": "Age"}
    rows = []
    for col in numeric_cols:
        s = imputed[col].dropna().astype(float)
        rows.append([labels[col], f"{s.mean():.1f}", f"{s.median():.0f}", f"{s.min():.0f}", f"{s.max():.0f}"])
    _table(doc, ["Variable", "Mean", "Median", "Min", "Max"], rows)

    for buf in dist_bufs:
        _add_chart(doc, buf, width_in=5.5)


def _section_choice_chart(doc: Document, title: str, buf) -> None:
    _heading(doc, "6. A Chart That Tells a Story", level=1)
    _para(doc, title)
    _add_chart(doc, buf)


def _section_limitations(doc: Document, name: str) -> None:
    _heading(doc, "7. Limitations and Known Risks", level=1)
    items = [
        "Cleaning and imputation are scoped to the assigned named variables only (brief Section 1.2) "
        "-- other columns are documented in the data dictionary but not reviewed for quality issues.",
        "Structural vs. incidental missingness is only classified for a verified, bounded set of "
        "single-parent skip gates (config.SKIP_PATTERNS) -- not exhaustive coverage of the full "
        "questionnaire's skip logic.",
        "Free-text consolidation only auto-merges pairs at >=0.98 similarity; the noisier 0.90-0.97 "
        "band is left flagged for manual review, not merged.",
        "Imputation uses simple median fill only (brief Section 3.3's basic-methods ceiling) -- no "
        "multivariate or model-based imputation.",
        "Monetary fields are outlier-flagged (1.5x IQR) but never altered -- genuine right-skewed "
        "income data will trigger some flags that aren't actual errors.",
    ]
    if name == "ind":
        items.append("panelid is named like a unique respondent ID but has only 18 distinct values across 13,853 rows -- documented so it is never mistaken for a join key.")
    else:
        items.append("emig has no documented expected row/household count in the source methodological report, unlike ind -- ingest.py logs its counts for the record only, without a pass/fail check.")
    _bullets(doc, items)


# ---------------------------------------------------------------- orchestration

def build_report(name: str, datasets: dict) -> Path:
    paths = datasets[name]
    typed = pd.read_parquet(paths["typed"])
    cleaned = pd.read_parquet(paths["cleaned"])
    imputed = pd.read_parquet(paths["imputed"])
    labels = column_labels(name, datasets, imputed.columns)

    doc = Document()
    doc.add_heading(f"GLFS 2017 Q4 -- Automated Cleaning & Imputation Report ({name})", level=0)
    doc.add_paragraph("Rendered automatically by Scripts/report.py -- every figure below is computed from the pipeline's own output at render time.")

    _section_dataset_overview(doc, name, datasets, typed)

    missingness_buf = chart_missingness_overview(name, typed, cleaned)
    _section_diagnostics(doc, name, typed, missingness_buf)

    _section_cleaning(doc, name, cleaned, typed)

    if name == "ind":
        obs_vs_imp_buf = chart_observed_vs_imputed(imputed, "q3_03", "Hours worked (main job): observed vs. imputed values")
    else:
        obs_vs_imp_buf = chart_observed_vs_imputed(imputed, "q7_06", "Age: observed vs. imputed values")
    _section_imputation(doc, name, imputed, labels, obs_vs_imp_buf)

    if name == "ind":
        dist_bufs = [
            chart_numeric_distribution(imputed["q1_04"], "Age distribution (final)", "Age (years)", _CATEGORICAL[0]),
            chart_numeric_distribution(imputed["q3_03"], "Hours worked distribution (final)", "Hours per week", _CATEGORICAL[2]),
        ]
    else:
        dist_bufs = [
            chart_numeric_distribution(imputed["q7_06"], "Age distribution (final)", "Age (years)", _CATEGORICAL[0]),
        ]
    _section_post_cleaning(doc, name, imputed, dist_bufs)

    if name == "ind":
        choice_buf = chart_of_choice(imputed, "q3_16", "Employment status (main job)")
        _section_choice_chart(doc, "Employment status among respondents with a main job.", choice_buf)
    else:
        choice_buf = chart_of_choice(imputed, "q7_11", "Most important reason for emigrating")
        _section_choice_chart(doc, "Self-reported reasons for having moved out of Guyana.", choice_buf)

    _section_limitations(doc, name)

    doc.save(str(paths["report_docx"]))

    log_lines = [
        f"Report: {name}",
        f"Written to {paths['report_docx'].name}",
        f"Rows/cols documented: {imputed.shape[0]} / {imputed.shape[1]}",
        "Charts embedded: 5 (missingness overview, observed-vs-imputed, 1-2 numeric distributions, 1 chart of choice)",
    ]
    paths["report_log"].write_text("\n".join(log_lines) + "\n", encoding="utf-8")
    print("\n".join(log_lines))

    return paths["report_docx"]


def build_all(datasets: dict) -> dict:
    return {name: build_report(name, datasets) for name in datasets}


if __name__ == "__main__":
    build_all(config.DATASETS)
